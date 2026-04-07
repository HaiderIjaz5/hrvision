import os
import re
import PyPDF2
import docx
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

def extract_text_from_file(filepath):
    """Extracts raw text from PDF or DOCX files."""
    text = ""
    try:
        ext = filepath.rsplit('.', 1)[1].lower()
        
        # --- Handle PDF Files ---
        if ext == 'pdf':
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
                        
        # --- Handle Word Documents ---
        elif ext in ['doc', 'docx']:
            doc = docx.Document(filepath)
            for para in doc.paragraphs:
                text += para.text + "\n"
                
    except Exception as e:
        print(f"⚠️ Error reading file {filepath}: {e}")
        
    return text.strip()

def clean_text(text):
    """Cleans text by making it lowercase and removing special characters."""
    if not text:
        return ""
    text = text.lower()
    # Remove everything except letters, numbers, and spaces
    text = re.sub(r'[^a-z0-9\s]', ' ', text)
    # Remove extra multiple spaces
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def calculate_resume_score(resume_path, job_description, mandatory_skills, candidate_skills):
    """
    Core AI scoring function.
    Compares the Resume text + Candidate manually entered skills
    against the Job Description text using Cosine Similarity.
    """
    
    # 1. Extract text from the candidate's uploaded resume file
    resume_raw_text = extract_text_from_file(resume_path)
    
    if not resume_raw_text and not candidate_skills:
        return 0 # If we got absolutely no data, score is 0
        
    # 2. Combine the resume text with the skills they typed into the form
    candidate_full_text = resume_raw_text + " " + " ".join(candidate_skills)
    
    # 3. Clean both texts to prepare them for Machine Learning math
    candidate_clean = clean_text(candidate_full_text)
    jd_clean = clean_text(job_description)

    if not jd_clean:
        return 0 # If the Admin left the JD blank, we can't score it

    # =======================================================
    # 4. THE NLP ENGINE: TF-IDF & Cosine Similarity
    # =======================================================
    documents = [jd_clean, candidate_clean]
    
    # TfidfVectorizer ignores common "stop words" like "and", "the", "is"
    vectorizer = TfidfVectorizer(stop_words='english')
    
    try:
        # Convert the texts into mathematical matrices
        tfidf_matrix = vectorizer.fit_transform(documents)
        
        # Calculate the angle (similarity) between the Job and the Resume
        similarity_matrix = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])
        
        # Convert the raw decimal (e.g., 0.75) into a percentage (75%)
        base_nlp_score = similarity_matrix[0][0] * 100
    except ValueError:
        # Failsafe: Happens if documents are empty or contain only stop words
        base_nlp_score = 0.0

    # =======================================================
    # 5. PENALTY SYSTEM: Mandatory Skills Check
    # =======================================================
    penalty = 0
    candidate_search_text = candidate_clean.lower()
    
    if mandatory_skills:
        missing_skills = 0
        for skill in mandatory_skills:
            if skill.lower() not in candidate_search_text:
                missing_skills += 1
        
        # Deduct 10 percentage points for EVERY mandatory skill they are missing
        penalty = missing_skills * 10

    # =======================================================
    # 6. FINAL CALCULATION
    # =======================================================
    final_score = int(base_nlp_score - penalty)
    
    # Ensure the score stays within a logical 0% to 100% boundary
    final_score = max(0, min(100, final_score))
    
    return final_score



    #new backup
    import os
import re
import PyPDF2
import docx

# --- NEW: TRUE DEEP LEARNING AI ---
from sentence_transformers import SentenceTransformer, util

# Load the pre-trained NLP model into memory.
# Note: The very first time you run your server, it will download a ~90MB model file from HuggingFace.
print("🧠 Loading Deep Learning AI Model (this might take a few seconds)...")
model = SentenceTransformer('all-MiniLM-L6-v2')
print("✅ AI Model successfully loaded and ready!")

def extract_text_from_file(filepath):
    """Extracts raw text from PDF or DOCX files."""
    text = ""
    try:
        ext = filepath.rsplit('.', 1)[1].lower()
        
        # --- Handle PDF Files ---
        if ext == 'pdf':
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
                        
        # --- Handle Word Documents ---
        elif ext in ['doc', 'docx']:
            doc = docx.Document(filepath)
            for para in doc.paragraphs:
                text += para.text + "\n"
                
    except Exception as e:
        print(f"⚠️ Error reading file {filepath}: {e}")
        
    return text.strip()

def clean_text(text):
    """Cleans text by making it lowercase and removing special characters."""
    if not text:
        return ""
    text = text.lower()
    # Remove everything except letters, numbers, and spaces
    text = re.sub(r'[^a-z0-9\s]', ' ', text)
    # Remove extra multiple spaces
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def calculate_resume_score(resume_path, job_description, mandatory_skills, candidate_skills):
    """
    ADVANCED AI SCORING:
    Uses BERT-based Sentence Transformers to understand the *semantic meaning* of the resume and compares it to the job description.
    """
    
    # 1. Extract text from the candidate's uploaded resume file
    resume_raw_text = extract_text_from_file(resume_path)
    
    if not resume_raw_text and not candidate_skills:
        return 0 
        
    # 2. Combine the resume text with the skills they typed into the form
    candidate_full_text = resume_raw_text + " " + " ".join(candidate_skills)
    
    # 3. Clean both texts
    candidate_clean = clean_text(candidate_full_text)
    jd_clean = clean_text(job_description)

    if not jd_clean:
        return 0 

    # =======================================================
    # 4. THE DEEP LEARNING ENGINE (Semantic Context Matching)
    # =======================================================
    try:
        # Convert the texts into high-dimensional numerical vectors (Embeddings)
        # This maps the *meaning* of the text into mathematical space
        candidate_embedding = model.encode(candidate_clean, convert_to_tensor=True)
        jd_embedding = model.encode(jd_clean, convert_to_tensor=True)
        
        # Calculate the Cosine Similarity between the two meaning-vectors
        cosine_scores = util.cos_sim(candidate_embedding, jd_embedding)
        
        # Extract the score and convert to a percentage
        base_ai_score = cosine_scores[0][0].item() * 100
    except Exception as e:
        print(f"AI Scoring Error: {e}")
        base_ai_score = 0.0

    # =======================================================
    # 5. PENALTY SYSTEM: Strict Requirements Check
    # =======================================================
    # Even if the AI thinks the context matches, HR needs exact skills.
    penalty = 0
    candidate_search_text = candidate_clean.lower()
    
    if mandatory_skills:
        missing_skills = 0
        for skill in mandatory_skills:
            if skill.lower() not in candidate_search_text:
                missing_skills += 1
        
        # Deduct 10 percentage points for EVERY mandatory skill they are missing
        penalty = missing_skills * 10

    # =======================================================
    # 6. FINAL CALCULATION
    # =======================================================
    final_score = int(base_ai_score - penalty)
    
    # Ensure the score stays within a logical 0% to 100% boundary
    final_score = max(0, min(100, final_score))
    
    return final_score